import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def generate_intuitive_docx(filename="Research_Overview_AttnRes_SiTU.docx"):
    doc = Document()

    # Set 0.75 inch margins for multi-page document
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Base style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)

    # Document Header Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run("Intuitive Guide: Modernizing Transformers with SSF-AttnRes & Kimi K3 SiTU-GLU")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1A, 0x2B, 0x4C)

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(8)
    run_sub = p_sub.add_run("A Conceptual & Technical Comparison: Traditional Architectures vs. Kimi K3 Innovations")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    # Divider
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(10)
    r_div = p_div.add_run("―" * 68)
    r_div.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
    r_div.font.bold = True

    # 1. Executive Overview & Intuitive Context
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(6)
    h1.paragraph_format.space_after = Pt(4)
    r1 = h1.add_run("1. Executive Overview & Intuitive Context")
    r1.font.size = Pt(13)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x1A, 0x2B, 0x4C)

    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(8)
    p1.paragraph_format.line_spacing = 1.15
    p1.add_run(
        "For the past decade, Large Language Models (LLMs) have relied on the standard Transformer architecture introduced by Vaswani et al. (2017). "
        "While standard Transformers have powered massive LLMs like GPT-4 and LLaMA, they suffer from fundamental representation bottlenecks "
        "and numerical instability in low-precision (FP16/FP8) environments. "
        "This project investigates two groundbreaking architectural alternatives: Sub-Sequence Free Attention Residuals (SSF-AttnRes) "
        "and Sigmoid Tanh Unit GLU (SiTU-GLU), recently introduced in Moonshot AI's Kimi K3 Technical Report."
    )

    # 2. What We Used Before (Traditional Transformer Foundations)
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(4)
    r2 = h2.add_run("2. What We Used Before: Traditional Transformer Bottlenecks")
    r2.font.size = Pt(13)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0x1A, 0x2B, 0x4C)

    # 2.1 Additive Residuals
    bp1 = doc.add_paragraph(style='List Bullet')
    bp1.paragraph_format.space_after = Pt(4)
    bp1.add_run("1. Pre-Norm Additive Residual Connections (x_{l+1} = x_l + f_l(x_l)): ").bold = True
    bp1.add_run(
        "In traditional Transformers, token representations travel down a strict sequential 'conveyor belt'. "
        "Each layer adds its transformation directly onto the running sum x_l. "
        "Intuitive Problem: Information from early token embeddings (y_0) becomes heavily diluted as network depth grows. "
        "Deep layers (e.g., Layer 24 or 32) cannot directly inspect pristine early features without wading through all intermediate transformations."
    )

    # 2.2 Unbounded SwiGLU
    bp2 = doc.add_paragraph(style='List Bullet')
    bp2.paragraph_format.space_after = Pt(8)
    bp2.add_run("2. Standard Un-Bounded SwiGLU Activations: ").bold = True
    bp2.add_run(
        "Modern LLMs use SwiGLU MLPs: SwiGLU(x) = (SiLU(xW_g) ⊙ xW_u)W_d. "
        "Intuitive Problem: SwiGLU outputs are completely un-bounded (-∞ to +∞). "
        "During low-precision training (FP16/FP8), high learning rates cause activation magnitudes to explode, "
        "leading to numerical overflow (NaNs), gradient norm spikes, and sudden loss divergence."
    )

    # 3. What Kimi K3 & SSF-AttnRes Introduced
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    r3 = h3.add_run("3. The New Paradigm: SSF-AttnRes & Kimi K3 SiTU-GLU")
    r3.font.size = Pt(13)
    r3.font.bold = True
    r3.font.color.rgb = RGBColor(0x1A, 0x2B, 0x4C)

    # 3.1 AttnRes
    bp3 = doc.add_paragraph(style='List Bullet')
    bp3.paragraph_format.space_after = Pt(4)
    bp3.add_run("1. Sub-Sequence Free Attention Residuals (SSF-AttnRes): ").bold = True
    bp3.add_run(
        "Instead of naive scalar addition, AttnRes equips each layer l with a learnable 'searchlight' (pseudo-query vector q_l). "
        "Layer l uses q_l to compute softmax attention over the token embedding y_0 and ALL prior layer outputs [y_0, y_1, ..., y_{l-1}]. "
        "Intuitive Breakthrough: Deep layers can dynamically route and retrieve exact features from any previous depth level, "
        "bypassing sequential dilution."
    )

    # 3.2 SiTU-GLU
    bp4 = doc.add_paragraph(style='List Bullet')
    bp4.paragraph_format.space_after = Pt(8)
    bp4.add_run("2. Sigmoid Tanh Unit GLU (SiTU-GLU from Kimi K3 Eq. 12): ").bold = True
    bp4.add_run(
        "Introduced in Moonshot AI's recent Kimi K3 Technical Report, SiTU-GLU applies smooth tanh softcapping to both gate and up branches: "
        "[β_1 tanh(xW_g / β_1) ⊙ Sigmoid(xW_g)] ⊙ [β_2 tanh(xW_u / β_2)] W_d. "
        "Intuitive Breakthrough: With parameters β_1 = 4.0 and β_2 = 25.0, SiTU-GLU places a strict mathematical ceiling on forward activations "
        "(||x||_∞ ≤ β_1 · β_2 = 100), eliminating activation magnitude explosion."
    )

    # 4. Controlled 4-Variant Research Matrix
    h4 = doc.add_paragraph()
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after = Pt(4)
    r4 = h4.add_run("4. Controlled 4-Variant Architectural Benchmark Matrix")
    r4.font.size = Pt(13)
    r4.font.bold = True
    r4.font.color.rgb = RGBColor(0x1A, 0x2B, 0x4C)

    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(6)
    p4.add_run("To isolate and evaluate these mechanisms, we benchmark a custom 50M LM on FineWeb-Edu across a 4-variant controlled matrix:")

    # Table
    table = doc.add_table(rows=5, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Variant #", "Architecture Name", "Core Mechanism & Paradigm Shift"]
    row_data = [
        ["Variant 0", "Baseline (Traditional)", "Pre-Norm Additive Residual (x + f(x)) + Standard Un-bounded SwiGLU MLP"],
        ["Variant 1", "+AttnRes (Depth Routing)", "AttnRes Dynamic Depth-Attention Routing + Standard SwiGLU MLP"],
        ["Variant 2", "+SiTU (Kimi K3 Bounded)", "Pre-Norm Additive Residual + Kimi K3 Bounded SiTU-GLU MLP (||x||_∞ ≤ 100)"],
        ["Variant 3", "+Both (AttnRes + SiTU)", "AttnRes Dynamic Depth Routing + Kimi K3 Bounded SiTU-GLU MLP"]
    ]

    col_widths = [Inches(1.0), Inches(2.0), Inches(4.0)]

    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_background(hdr_cells[i], "1A2B4C")
        set_cell_margins(hdr_cells[i], top=80, bottom=80)

    for r_idx, r_vals in enumerate(row_data):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(r_vals):
            row_cells[c_idx].text = val
            p = row_cells[c_idx].paragraphs[0]
            p.runs[0].font.size = Pt(9)
            set_cell_margins(row_cells[c_idx], top=60, bottom=60)
            if r_idx % 2 == 1:
                set_cell_background(row_cells[c_idx], "F7FAFC")

    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    # 5. Primary Research Questions
    h5 = doc.add_paragraph()
    h5.paragraph_format.space_before = Pt(10)
    h5.paragraph_format.space_after = Pt(4)
    r5 = h5.add_run("5. Core Research Objectives & Hypotheses")
    r5.font.size = Pt(13)
    r5.font.bold = True
    r5.font.color.rgb = RGBColor(0x1A, 0x2B, 0x4C)

    q1 = doc.add_paragraph(style='List Bullet')
    q1.paragraph_format.space_after = Pt(3)
    q1.add_run("RQ1 (Representation Retrieval): ").bold = True
    q1.add_run("Does depth-attention routing (AttnRes) allow deep layers to shortcut representation retrieval, boosting token sample efficiency?")

    q2 = doc.add_paragraph(style='List Bullet')
    q2.paragraph_format.space_after = Pt(3)
    q2.add_run("RQ2 (Forward vs. Backward Bounding): ").bold = True
    q2.add_run("Does capping forward activations (||x||_∞ ≤ 100 via SiTU-GLU) automatically protect against backward gradient norm explosions?")

    q3 = doc.add_paragraph(style='List Bullet')
    q3.paragraph_format.space_after = Pt(6)
    q3.add_run("RQ3 (Un-Normalized Stress Test): ").bold = True
    q3.add_run("How do these architectural variants perform when external normalization 'safety nets' (e.g. sub-block RMSNorm) are intentionally omitted under un-normalized residual stress?")

    # Footer
    p_foot = doc.add_paragraph()
    p_foot.paragraph_format.space_before = Pt(10)
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot = p_foot.add_run("SSF-AttnRes Research Framework | Custom Transformer Architecture Investigation")
    r_foot.font.size = Pt(8.5)
    r_foot.font.italic = True
    r_foot.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)

    doc.save(filename)
    print(f"Successfully generated DOCX: {filename}")


def generate_intuitive_pdf(filename="Research_Overview_Intuitive_Guide.pdf"):
    doc = SimpleDocTemplate(
        filename,
        pagesize=letter,
        leftMargin=30,
        rightMargin=30,
        topMargin=30,
        bottomMargin=30
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=15,
        leading=18,
        textColor=colors.HexColor("#1A2B4C"),
        spaceAfter=2
    )

    subtitle_style = ParagraphStyle(
        'DocSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=9.5,
        leading=12.5,
        textColor=colors.HexColor("#4A5568"),
        spaceAfter=6
    )

    section_heading = ParagraphStyle(
        'SectionHeading',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#1A2B4C"),
        spaceBefore=6,
        spaceAfter=2
    )

    body_style = ParagraphStyle(
        'BodyTextCustom',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8.5,
        leading=11.5,
        textColor=colors.HexColor("#2D3748"),
        spaceAfter=4
    )

    bullet_style = ParagraphStyle(
        'BulletCustom',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8.2,
        leading=11.2,
        textColor=colors.HexColor("#2D3748"),
        leftIndent=12,
        firstLineIndent=-9,
        spaceAfter=3
    )

    table_header_style = ParagraphStyle(
        'TableHeader',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8.5,
        leading=11,
        textColor=colors.white,
        alignment=TA_CENTER
    )

    table_cell_style = ParagraphStyle(
        'TableCell',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8,
        leading=10.5,
        textColor=colors.HexColor("#1A202C"),
        alignment=TA_CENTER
    )

    table_cell_left = ParagraphStyle(
        'TableCellLeft',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8,
        leading=10.5,
        textColor=colors.HexColor("#1A202C"),
        alignment=TA_LEFT
    )

    elements = []

    # Title & Subtitle
    elements.append(Paragraph("Intuitive Guide: Modernizing Transformers with SSF-AttnRes & Kimi K3 SiTU-GLU", title_style))
    elements.append(Paragraph("A Conceptual & Technical Comparison: Traditional Architectures vs. Kimi K3 Innovations", subtitle_style))
    elements.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#2B6CB0"), spaceAfter=5))

    # 1. Executive Overview & Intuitive Context
    elements.append(Paragraph("1. Executive Overview & Intuitive Context", section_heading))
    bg_text = (
        "For the past decade, Large Language Models (LLMs) have relied on the standard Transformer architecture "
        "introduced by Vaswani et al. (2017). While standard Transformers power massive LLMs like GPT-4 and LLaMA, "
        "they suffer from fundamental representation bottlenecks and numerical instability in low-precision (FP16/FP8) environments. "
        "This project investigates two groundbreaking architectural alternatives: <b>Sub-Sequence Free Attention Residuals (SSF-AttnRes)</b> "
        "and <b>Sigmoid Tanh Unit GLU (SiTU-GLU)</b>, recently introduced in Moonshot AI's Kimi K3 Technical Report."
    )
    elements.append(Paragraph(bg_text, body_style))

    # 2. What We Used Before (Traditional Transformer Foundations)
    elements.append(Paragraph("2. What We Used Before: Traditional Transformer Bottlenecks", section_heading))
    b1 = "• <b>1. Pre-Norm Additive Residual Connections (<i>x<sub>l+1</sub> = x<sub>l</sub> + f<sub>l</sub>(x<sub>l</sub>)</i>)</b>: In traditional Transformers, token representations travel down a strict sequential 'conveyor belt'. Each layer adds its transformation directly onto the running sum. <i>Intuitive Problem</i>: Information from early token embeddings (<i>y<sub>0</sub></i>) becomes heavily diluted as network depth grows. Deep layers cannot directly inspect pristine early features without wading through all intermediate transformations."
    b2 = "• <b>2. Standard Un-Bounded SwiGLU Activations</b>: Modern LLMs use SwiGLU MLPs: <i>SwiGLU(x) = (SiLU(xW<sub>g</sub>) · xW<sub>u</sub>)W<sub>d</sub></i>. <i>Intuitive Problem</i>: SwiGLU outputs are completely un-bounded (-∞ to +∞). During low-precision training (FP16/FP8), high learning rates cause activation magnitudes to explode, leading to numerical overflow (NaNs), gradient norm spikes, and loss divergence."
    elements.append(Paragraph(b1, bullet_style))
    elements.append(Paragraph(b2, bullet_style))

    # 3. What Kimi K3 & SSF-AttnRes Introduced
    elements.append(Paragraph("3. The New Paradigm: SSF-AttnRes & Kimi K3 SiTU-GLU", section_heading))
    n1 = "• <b>1. Sub-Sequence Free Attention Residuals (SSF-AttnRes)</b>: Instead of naive scalar addition, AttnRes equips each layer <i>l</i> with a learnable 'searchlight' (pseudo-query vector <i>q<sub>l</sub></i>). Layer <i>l</i> uses <i>q<sub>l</sub></i> to compute softmax attention over the token embedding <i>y<sub>0</sub></i> and ALL prior layer outputs [<i>y<sub>0</sub>, y<sub>1</sub>, ..., y<sub>l-1</sub></i>]. <i>Intuitive Breakthrough</i>: Deep layers can dynamically route and retrieve exact features from any previous depth level, bypassing sequential dilution."
    n2 = "• <b>2. Sigmoid Tanh Unit GLU (SiTU-GLU from Kimi K3 Eq. 12)</b>: Introduced in Moonshot AI's recent Kimi K3 Technical Report, SiTU-GLU applies smooth tanh softcapping to both gate and up branches: <i>[β<sub>1</sub> tanh(xW<sub>g</sub> / β<sub>1</sub>) · Sigmoid(xW<sub>g</sub>)] · [β<sub>2</sub> tanh(xW<sub>u</sub> / β<sub>2</sub>)] W<sub>d</sub></i>. <i>Intuitive Breakthrough</i>: With parameters β<sub>1</sub> = 4.0 and β<sub>2</sub> = 25.0, SiTU-GLU places a strict mathematical ceiling on forward activations (||x||<sub>∞</sub> ≤ β<sub>1</sub> · β<sub>2</sub> = 100), eliminating activation magnitude explosion."
    elements.append(Paragraph(n1, bullet_style))
    elements.append(Paragraph(n2, bullet_style))

    # 4. Controlled 4-Variant Benchmark Design
    elements.append(Paragraph("4. Controlled 4-Variant Research Matrix", section_heading))
    elements.append(Paragraph("To isolate and evaluate these mechanisms, we benchmark a custom 50M LM on FineWeb-Edu across a 4-variant controlled matrix:", body_style))

    table_data = [
        [
            Paragraph("Variant #", table_header_style),
            Paragraph("Architecture Name", table_header_style),
            Paragraph("Core Mechanism & Paradigm Shift", table_header_style),
        ],
        [
            Paragraph("Variant 0", table_cell_style),
            Paragraph("Baseline (Traditional)", table_cell_left),
            Paragraph("Pre-Norm Additive Residual (x + f(x)) + Standard Un-bounded SwiGLU MLP", table_cell_left),
        ],
        [
            Paragraph("Variant 1", table_cell_style),
            Paragraph("+AttnRes (Depth Routing)", table_cell_left),
            Paragraph("AttnRes Dynamic Depth-Attention Routing + Standard SwiGLU MLP", table_cell_left),
        ],
        [
            Paragraph("Variant 2", table_cell_style),
            Paragraph("+SiTU (Kimi K3 Bounded)", table_cell_left),
            Paragraph("Pre-Norm Additive Residual + Kimi K3 Bounded SiTU-GLU MLP (||x||<sub>∞</sub> ≤ 100)", table_cell_left),
        ],
        [
            Paragraph("Variant 3", table_cell_style),
            Paragraph("+Both (AttnRes + SiTU)", table_cell_left),
            Paragraph("AttnRes Dynamic Depth Routing + Kimi K3 Bounded SiTU-GLU MLP", table_cell_left),
        ]
    ]

    t = Table(table_data, colWidths=[60, 130, 360])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#1A2B4C")),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E0")),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7FAFC")]),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 3))

    # 5. Core Research Objectives
    elements.append(Paragraph("5. Core Research Objectives & Hypotheses", section_heading))
    q1 = "• <b>RQ1 (Representation Retrieval)</b>: Does depth-attention routing (AttnRes) allow deep layers to shortcut representation retrieval, boosting token sample efficiency?"
    q2 = "• <b>RQ2 (Forward vs. Backward Bounding)</b>: Does capping forward activations (||x||<sub>∞</sub> ≤ 100 via SiTU-GLU) automatically protect against backward gradient norm explosions?"
    q3 = "• <b>RQ3 (Un-Normalized Stress Test)</b>: How do these architectural variants perform when external normalization 'safety nets' (e.g. sub-block RMSNorm) are intentionally omitted under un-normalized residual stress?"
    elements.append(Paragraph(q1, bullet_style))
    elements.append(Paragraph(q2, bullet_style))
    elements.append(Paragraph(q3, bullet_style))

    # Footer
    elements.append(HRFlowable(width="100%", thickness=0.75, color=colors.HexColor("#E2E8F0"), spaceBefore=4, spaceAfter=2))
    elements.append(Paragraph("SSF-AttnRes Research Framework | Custom Transformer Architecture Investigation", ParagraphStyle('Footer', parent=styles['Normal'], fontName='Helvetica-Oblique', fontSize=7.5, textColor=colors.HexColor("#A0AEC0"), alignment=TA_CENTER)))

    doc.build(elements)
    print(f"Successfully generated intuitive PDF: {filename}")

if __name__ == "__main__":
    try:
        generate_intuitive_docx("Research_Overview_Intuitive_Guide.docx")
    except PermissionError:
        generate_intuitive_docx("Research_Overview_Intuitive_Guide_v2.docx")

    try:
        generate_intuitive_pdf("Research_Overview_Intuitive_Guide.pdf")
    except PermissionError:
        generate_intuitive_pdf("Research_Overview_Intuitive_Guide_v2.pdf")

    try:
        generate_intuitive_pdf("Research_Summary_OnePage.pdf")
    except PermissionError:
        pass
