import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def generate_docx(filename="Research_Overview_AttnRes_SiTU.docx"):
    doc = Document()

    # Set 0.5 inch margins for 1-page fit
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)

    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run("Research Brief: Custom Transformer Architecture with SSF-AttnRes & SiTU-GLU")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1A, 0x2B, 0x4C)

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(6)
    run_sub = p_sub.add_run("Focus: Evaluating Dynamic Depth Routing & Bounded Activations (Kimi K3) | Model: 50M LM | Dataset: FineWeb-Edu")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(9.5)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    # Divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(6)
    r_div = p_div.add_run("―" * 65)
    r_div.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
    r_div.font.bold = True

    # 1. Background & Motivation
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(4)
    h1.paragraph_format.space_after = Pt(3)
    r1 = h1.add_run("1. Background & Motivation")
    r1.font.size = Pt(11.5)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x1A, 0x2B, 0x4C)

    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(5)
    p1.paragraph_format.line_spacing = 1.15
    p1.add_run(
        "Standard Transformer architectures (e.g., LLaMA, GPT) rely on static additive residual connections (x + f(x)) and un-bounded SwiGLU activations. "
        "Additive residual pathways accumulate features sequentially, limiting direct representation routing from early to deep layers. "
        "Furthermore, un-bounded activations in standard SwiGLU can experience dynamic range explosion and loss instability under low-precision training (FP16/FP8). "
        "This research explores custom architectural modifications to overcome these fundamental bottlenecks."
    )

    # 2. Core Architectural Innovations
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(4)
    h2.paragraph_format.space_after = Pt(3)
    r2 = h2.add_run("2. Core Architectural Innovations")
    r2.font.size = Pt(11.5)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0x1A, 0x2B, 0x4C)

    bp1 = doc.add_paragraph(style='List Bullet')
    bp1.paragraph_format.space_after = Pt(3)
    r_b1_title = bp1.add_run("Sub-Sequence Free Attention Residuals (SSF-AttnRes): ")
    r_b1_title.bold = True
    bp1.add_run(
        "Replaces scalar addition with dynamic depth-attention routing. Layer l utilizes a learnable pseudo-query vector q_l "
        "to compute softmax attention over the token embedding y_0 and all preceding layer representations [y_0, y_1, ..., y_{l-1}], "
        "enabling selective information extraction across depth."
    )

    bp2 = doc.add_paragraph(style='List Bullet')
    bp2.paragraph_format.space_after = Pt(5)
    r_b2_title = bp2.add_run("Sigmoid Tanh Unit GLU (SiTU-GLU from Kimi K3): ")
    r_b2_title.bold = True
    bp2.add_run(
        "Introduced in the recent Kimi K3 technical report (Eq. 12), SiTU-GLU incorporates smooth softcapping via tanh functions "
        "on both the gate and up branches: [β_1 tanh(xW_g / β_1) ⊙ Sigmoid(xW_g)] ⊙ [β_2 tanh(xW_u / β_2)] W_d. "
        "This strictly guarantees bounded forward activation magnitude (||x||_∞ ≤ β_1 · β_2 = 100)."
    )

    # 3. Controlled 4-Variant Benchmark Design
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(4)
    h3.paragraph_format.space_after = Pt(3)
    r3 = h3.add_run("3. Controlled 4-Variant Research Matrix")
    r3.font.size = Pt(11.5)
    r3.font.bold = True
    r3.font.color.rgb = RGBColor(0x1A, 0x2B, 0x4C)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(4)
    p3.add_run("To isolate and evaluate the individual and combined impact of these innovations, we construct a 4-variant matrix across 3 independent random seeds (12 experiments total, trained on FineWeb-Edu):")

    # Table
    table = doc.add_table(rows=5, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Variant #", "Architecture Name", "Key Architectural Mechanism"]
    row_data = [
        ["0", "Baseline", "Standard Pre-Norm Additive Residual + SwiGLU MLP"],
        ["1", "Variant 1: +AttnRes", "AttnRes Depth-Attention Routing + SwiGLU MLP"],
        ["2", "Variant 2: +SiTU", "Pre-Norm Additive Residual + Bounded SiTU-GLU MLP (Kimi K3)"],
        ["3", "Variant 3: +Both", "AttnRes Depth Routing + Bounded SiTU-GLU MLP"]
    ]

    col_widths = [Inches(0.8), Inches(1.8), Inches(4.4)]

    # Format header
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_background(hdr_cells[i], "1A2B4C")
        set_cell_margins(hdr_cells[i], top=80, bottom=80)

    # Format rows
    for r_idx, r_vals in enumerate(row_data):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(r_vals):
            row_cells[c_idx].text = val
            p = row_cells[c_idx].paragraphs[0]
            p.runs[0].font.size = Pt(8.5)
            set_cell_margins(row_cells[c_idx], top=60, bottom=60)
            if r_idx % 2 == 1:
                set_cell_background(row_cells[c_idx], "F7FAFC")

    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    # 4. Research Objectives & Key Questions
    h4 = doc.add_paragraph()
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after = Pt(3)
    r4 = h4.add_run("4. Primary Research Objectives")
    r4.font.size = Pt(11.5)
    r4.font.bold = True
    r4.font.color.rgb = RGBColor(0x1A, 0x2B, 0x4C)

    q1 = doc.add_paragraph(style='List Bullet')
    q1.paragraph_format.space_after = Pt(2)
    q1.add_run("RQ1 (Representation Routing): ").bold = True
    q1.add_run("Does depth-attention routing accelerate language modeling sample efficiency and improve peak validation loss?")

    q2 = doc.add_paragraph(style='List Bullet')
    q2.paragraph_format.space_after = Pt(2)
    q2.add_run("RQ2 (Activation & Gradient Bounding): ").bold = True
    q2.add_run("How effectively does SiTU-GLU activation softcapping suppress loss spikes and stabilize FP16 training dynamic range?")

    q3 = doc.add_paragraph(style='List Bullet')
    q3.paragraph_format.space_after = Pt(4)
    q3.add_run("RQ3 (Architectural Coupling): ").bold = True
    q3.add_run("Does combining depth routing with activation softcapping produce synergistic benefits for deep Transformer scaling?")

    # Footer note
    p_foot = doc.add_paragraph()
    p_foot.paragraph_format.space_before = Pt(6)
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot = p_foot.add_run("SSF-AttnRes Research Framework | Custom Transformer Architecture Investigation")
    r_foot.font.size = Pt(8)
    r_foot.font.italic = True
    r_foot.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)

    doc.save(filename)
    print(f"Successfully generated DOCX: {filename}")


def generate_pdf(filename="Research_Overview_AttnRes_SiTU.pdf"):
    doc = SimpleDocTemplate(
        filename,
        pagesize=letter,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=16,
        leading=20,
        textColor=colors.HexColor("#1A2B4C"),
        spaceAfter=3
    )

    subtitle_style = ParagraphStyle(
        'DocSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor("#4A5568"),
        spaceAfter=8
    )

    section_heading = ParagraphStyle(
        'SectionHeading',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11.5,
        leading=15,
        textColor=colors.HexColor("#1A2B4C"),
        spaceBefore=7,
        spaceAfter=3
    )

    body_style = ParagraphStyle(
        'BodyTextCustom',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12.5,
        textColor=colors.HexColor("#2D3748"),
        spaceAfter=4
    )

    bullet_style = ParagraphStyle(
        'BulletCustom',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8.5,
        leading=12,
        textColor=colors.HexColor("#2D3748"),
        leftIndent=12,
        firstLineIndent=-10,
        spaceAfter=3
    )

    table_header_style = ParagraphStyle(
        'TableHeader',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8.5,
        leading=11,
        textColor=colors.white,
        alignment=TA_CENTER
    )

    table_cell_style = ParagraphStyle(
        'TableCell',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8,
        leading=10.5,
        textColor=colors.HexColor("#1A202C"),
        alignment=TA_CENTER
    )

    table_cell_left = ParagraphStyle(
        'TableCellLeft',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8,
        leading=10.5,
        textColor=colors.HexColor("#1A202C"),
        alignment=TA_LEFT
    )

    elements = []

    # Title & Subtitle
    elements.append(Paragraph("Research Brief: Custom Transformer Architecture with SSF-AttnRes & SiTU-GLU", title_style))
    elements.append(Paragraph("Focus: Dynamic Depth Routing & Bounded Activations (Kimi K3) | Model: 50M LM | Dataset: FineWeb-Edu", subtitle_style))
    elements.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#2B6CB0"), spaceAfter=6))

    # 1. Background & Motivation
    elements.append(Paragraph("1. Background & Motivation", section_heading))
    bg_text = (
        "Standard Transformer architectures (e.g., LLaMA, GPT) rely on static additive residual connections "
        "(<i>x + f(x)</i>) and un-bounded SwiGLU activations. Additive residual pathways accumulate features "
        "sequentially, limiting direct representation routing from early to deep layers. Furthermore, un-bounded "
        "activations in standard SwiGLU can experience dynamic range explosion and loss instability under low-precision "
        "training (FP16/FP8). This research explores custom architectural modifications to overcome these fundamental bottlenecks."
    )
    elements.append(Paragraph(bg_text, body_style))

    # 2. Core Architectural Innovations
    elements.append(Paragraph("2. Core Architectural Innovations", section_heading))
    b1 = "• <b>Sub-Sequence Free Attention Residuals (SSF-AttnRes)</b>: Replaces scalar addition with dynamic depth-attention routing. Layer <i>l</i> utilizes a learnable pseudo-query vector <i>q<sub>l</sub></i> to compute softmax attention over the token embedding <i>y<sub>0</sub></i> and all preceding layer representations [<i>y<sub>0</sub>, y<sub>1</sub>, ..., y<sub>l-1</sub></i>], enabling selective information extraction across depth."
    b2 = "• <b>Sigmoid Tanh Unit GLU (SiTU-GLU from Kimi K3)</b>: Introduced in the recent Kimi K3 technical report (Eq. 12), SiTU-GLU incorporates smooth softcapping via tanh functions on both gate and up branches: [β<sub>1</sub> tanh(xW<sub>g</sub> / β<sub>1</sub>) · Sigmoid(xW<sub>g</sub>)] · [β<sub>2</sub> tanh(xW<sub>u</sub> / β<sub>2</sub>)] W<sub>d</sub>. This strictly guarantees bounded forward activation magnitude (||x||<sub>∞</sub> ≤ β<sub>1</sub> · β<sub>2</sub> = 100)."
    elements.append(Paragraph(b1, bullet_style))
    elements.append(Paragraph(b2, bullet_style))

    # 3. 4-Variant Benchmark Design
    elements.append(Paragraph("3. Controlled 4-Variant Research Matrix", section_heading))
    elements.append(Paragraph("To isolate and evaluate the individual and combined impact of these innovations, we construct a 4-variant matrix across 3 independent random seeds (12 experiments total, 1B tokens each):", body_style))

    table_data = [
        [
            Paragraph("Variant #", table_header_style),
            Paragraph("Architecture Name", table_header_style),
            Paragraph("Key Architectural Mechanism", table_header_style),
        ],
        [
            Paragraph("0", table_cell_style),
            Paragraph("Baseline", table_cell_left),
            Paragraph("Standard Pre-Norm Additive Residual + SwiGLU MLP", table_cell_left),
        ],
        [
            Paragraph("1", table_cell_style),
            Paragraph("Variant 1: +AttnRes", table_cell_left),
            Paragraph("AttnRes Depth-Attention Routing + SwiGLU MLP", table_cell_left),
        ],
        [
            Paragraph("2", table_cell_style),
            Paragraph("Variant 2: +SiTU", table_cell_left),
            Paragraph("Pre-Norm Additive Residual + Bounded SiTU-GLU MLP (Kimi K3)", table_cell_left),
        ],
        [
            Paragraph("3", table_cell_style),
            Paragraph("Variant 3: +Both", table_cell_left),
            Paragraph("AttnRes Depth Routing + Bounded SiTU-GLU MLP", table_cell_left),
        ]
    ]

    t = Table(table_data, colWidths=[55, 125, 330])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#1A2B4C")),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E0")),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7FAFC")]),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
    ]))
    elements.append(t)

    # 4. Primary Research Objectives
    elements.append(Paragraph("4. Primary Research Objectives", section_heading))
    q1 = "• <b>RQ1 (Representation Routing)</b>: Does depth-attention routing accelerate language modeling sample efficiency and improve peak validation loss?"
    q2 = "• <b>RQ2 (Activation & Gradient Bounding)</b>: How effectively does SiTU-GLU activation softcapping suppress loss spikes and stabilize FP16 training dynamic range?"
    q3 = "• <b>RQ3 (Architectural Coupling)</b>: Does combining depth routing with activation softcapping produce synergistic benefits for deep Transformer scaling?"
    elements.append(Paragraph(q1, bullet_style))
    elements.append(Paragraph(q2, bullet_style))
    elements.append(Paragraph(q3, bullet_style))

    # Footer
    elements.append(HRFlowable(width="100%", thickness=0.75, color=colors.HexColor("#E2E8F0"), spaceBefore=6, spaceAfter=4))
    elements.append(Paragraph("SSF-AttnRes Research Framework | Custom Transformer Architecture Investigation", ParagraphStyle('Footer', parent=styles['Normal'], fontName='Helvetica-Oblique', fontSize=7.5, textColor=colors.HexColor("#A0AEC0"), alignment=TA_CENTER)))

    doc.build(elements)
    print(f"Successfully generated PDF: {filename}")

if __name__ == "__main__":
    generate_docx("Research_Overview_AttnRes_SiTU.docx")
    generate_pdf("Research_Overview_AttnRes_SiTU.pdf")
    # Also overwrite Research_Summary_OnePage.pdf so both names have the clean idea version
    generate_pdf("Research_Summary_OnePage.pdf")
