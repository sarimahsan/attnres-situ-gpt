import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, HRFlowable
)

PRIMARY_HEX = "1E3A8A"     # Deep Blue
SECONDARY_HEX = "2563EB"   # Royal Blue
DARK_TEXT_HEX = "1F2937"   # Dark Gray
LIGHT_BG_HEX = "F3F4F6"    # Soft Light Gray
BORDER_HEX = "E5E7EB"      # Border Gray

PRIMARY_COLOR = colors.HexColor(f"#{PRIMARY_HEX}")
SECONDARY_COLOR = colors.HexColor(f"#{SECONDARY_HEX}")
DARK_TEXT_COLOR = colors.HexColor(f"#{DARK_TEXT_HEX}")
LIGHT_BG_COLOR = colors.HexColor(f"#{LIGHT_BG_HEX}")
BORDER_COLOR = colors.HexColor(f"#{BORDER_HEX}")

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def clean_md_formatting(text):
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    return text

def parse_markdown_to_docx(md_paths, docx_path):
    if isinstance(md_paths, str):
        md_paths = [md_paths]

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    for idx, md_path in enumerate(md_paths):
        if idx > 0:
            doc.add_page_break()

        with open(md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        in_code_block = False
        code_lines = []
        in_table = False
        table_lines = []

        for line in lines:
            raw_line = line.rstrip('\n')
            stripped = raw_line.strip()

            if stripped.startswith('```'):
                if in_code_block:
                    code_text = '\n'.join(code_lines)
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(8)
                    p.paragraph_format.left_indent = Inches(0.2)

                    table = doc.add_table(rows=1, cols=1)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    cell = table.cell(0, 0)
                    set_cell_background(cell, LIGHT_BG_HEX)
                    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)

                    cp = cell.paragraphs[0]
                    cp.paragraph_format.space_before = Pt(4)
                    cp.paragraph_format.space_after = Pt(4)
                    run = cp.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                    code_lines = []
                continue

            if in_code_block:
                code_lines.append(raw_line)
                continue

            if '|' in stripped and ('---' in stripped or len(stripped.split('|')) > 2):
                table_lines.append(stripped)
                in_table = True
                continue
            elif in_table:
                process_docx_table(doc, table_lines)
                table_lines = []
                in_table = False

            if not stripped:
                continue

            if stripped.startswith('# '):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(16)
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(clean_md_formatting(stripped[2:]))
                run.font.name = 'Segoe UI'
                run.font.size = Pt(20)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
            elif stripped.startswith('## '):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(clean_md_formatting(stripped[3:]))
                run.font.name = 'Segoe UI'
                run.font.size = Pt(15)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            elif stripped.startswith('### '):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(clean_md_formatting(stripped[4:]))
                run.font.name = 'Segoe UI'
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
            elif stripped.startswith('- ') or stripped.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(3)
                add_formatted_runs(p, clean_md_formatting(stripped[2:]))
            elif re.match(r'^\d+\.\s', stripped):
                content = re.sub(r'^\d+\.\s', '', stripped)
                p = doc.add_paragraph(style='List Number')
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(3)
                add_formatted_runs(p, clean_md_formatting(content))
            elif stripped.startswith('> '):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run(clean_md_formatting(stripped[2:]))
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(5)
                p.paragraph_format.line_spacing = 1.15
                add_formatted_runs(p, clean_md_formatting(stripped))

        if in_table and table_lines:
            process_docx_table(doc, table_lines)

    doc.save(docx_path)
    print(f"Generated DOCX: {docx_path}")

def add_formatted_runs(paragraph, text):
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
            run.font.name = 'Calibri'
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        else:
            run = paragraph.add_run(part)
            run.font.name = 'Calibri'

def process_docx_table(doc, lines):
    rows_data = []
    for line in lines:
        if '---' in line:
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            rows_data.append(cells)

    if not rows_data:
        return

    num_rows = len(rows_data)
    num_cols = max(len(r) for r in rows_data)

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row in enumerate(rows_data):
        for c_idx, cell_text in enumerate(row):
            if c_idx < num_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = clean_md_formatting(cell_text)

                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)

                if r_idx == 0:
                    set_cell_background(cell, PRIMARY_HEX)
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.name = 'Calibri'
                else:
                    if r_idx % 2 == 1:
                        set_cell_background(cell, "F9FAFB")
                    else:
                        set_cell_background(cell, "FFFFFF")
                    for run in p.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(9.5)

def parse_markdown_to_pdf(md_paths, pdf_path):
    if isinstance(md_paths, str):
        md_paths = [md_paths]

    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=letter,
        leftMargin=40,
        rightMargin=40,
        topMargin=40,
        bottomMargin=40
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        textColor=PRIMARY_COLOR,
        spaceAfter=12
    )

    h1_style = ParagraphStyle(
        'DocH1',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=18,
        textColor=PRIMARY_COLOR,
        spaceBefore=14,
        spaceAfter=8,
        keepWithNext=True
    )

    h2_style = ParagraphStyle(
        'DocH2',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=11.5,
        leading=15,
        textColor=SECONDARY_COLOR,
        spaceBefore=10,
        spaceAfter=6,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        textColor=DARK_TEXT_COLOR,
        spaceAfter=6
    )

    bullet_style = ParagraphStyle(
        'DocBullet',
        parent=body_style,
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=4
    )

    code_style = ParagraphStyle(
        'DocCode',
        parent=styles['Code'],
        fontName='Courier',
        fontSize=8.5,
        leading=11,
        textColor=DARK_TEXT_COLOR,
        backColor=LIGHT_BG_COLOR,
        spaceBefore=4,
        spaceAfter=6,
        leftIndent=10,
        rightIndent=10
    )

    story = []

    for idx, md_path in enumerate(md_paths):
        if idx > 0:
            story.append(PageBreak())

        with open(md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        in_code_block = False
        code_lines = []
        in_table = False
        table_lines = []

        for line in lines:
            raw_line = line.rstrip('\n')
            stripped = raw_line.strip()

            if stripped.startswith('```'):
                if in_code_block:
                    code_text = '<br/>'.join([c.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace(' ', '&nbsp;') for c in code_lines])
                    story.append(Paragraph(code_text, code_style))
                    story.append(Spacer(1, 4))
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                    code_lines = []
                continue

            if in_code_block:
                code_lines.append(raw_line)
                continue

            if '|' in stripped and ('---' in stripped or len(stripped.split('|')) > 2):
                table_lines.append(stripped)
                in_table = True
                continue
            elif in_table:
                process_pdf_table(story, table_lines, body_style)
                table_lines = []
                in_table = False

            if not stripped:
                continue

            formatted_text = format_md_to_reportlab(stripped)

            if stripped.startswith('# '):
                story.append(Paragraph(formatted_text[2:], title_style))
                story.append(HRFlowable(width="100%", thickness=1.5, color=PRIMARY_COLOR, spaceAfter=10))
            elif stripped.startswith('## '):
                story.append(Paragraph(formatted_text[3:], h1_style))
            elif stripped.startswith('### '):
                story.append(Paragraph(formatted_text[4:], h2_style))
            elif stripped.startswith('- ') or stripped.startswith('* '):
                story.append(Paragraph(f"• {formatted_text[2:]}", bullet_style))
            elif re.match(r'^\d+\.\s', stripped):
                story.append(Paragraph(formatted_text, bullet_style))
            else:
                story.append(Paragraph(formatted_text, body_style))

        if in_table and table_lines:
            process_pdf_table(story, table_lines, body_style)

    doc.build(story)
    print(f"Generated PDF: {pdf_path}")

def format_md_to_reportlab(text):
    text = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'`([^`]+)`', r'<font name="Courier" size="8.5" color="#1F2937"><b>\1</b></font>', text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    return text

def process_pdf_table(story, lines, body_style):
    rows_data = []
    for line in lines:
        if '---' in line:
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            rows_data.append(cells)

    if not rows_data:
        return

    table_data = []
    for r_idx, row in enumerate(rows_data):
        row_cells = []
        for cell_text in row:
            fmt = format_md_to_reportlab(cell_text)
            if r_idx == 0:
                p_style = ParagraphStyle(
                    'TH', parent=body_style, fontName='Helvetica-Bold', textColor=colors.white, fontSize=9, leading=12
                )
            else:
                p_style = ParagraphStyle(
                    'TD', parent=body_style, fontSize=8.5, leading=11
                )
            row_cells.append(Paragraph(fmt, p_style))
        table_data.append(row_cells)

    pdf_table = Table(table_data)
    ts = [
        ('BACKGROUND', (0, 0), (-1, 0), PRIMARY_COLOR),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('GRID', (0, 0), (-1, -1), 0.5, BORDER_COLOR),
    ]

    for r in range(1, len(table_data)):
        if r % 2 == 1:
            ts.append(('BACKGROUND', (0, r), (-1, r), colors.HexColor('#F9FAFB')))

    pdf_table.setStyle(TableStyle(ts))
    story.append(Spacer(1, 4))
    story.append(pdf_table)
    story.append(Spacer(1, 8))

def main():
    docs_dir = "docs"
    files = [
        "01_MAIN_IDEA.md",
        "02_ARCHITECTURE_DETAILS.md",
        "03_EXPERIMENTS_AND_SCHEDULE.md"
    ]

    full_paths = []
    for f in files:
        md_path = os.path.join(docs_dir, f)
        base_name = os.path.splitext(f)[0]
        docx_path = os.path.join(docs_dir, f"{base_name}.docx")
        pdf_path = os.path.join(docs_dir, f"{base_name}.pdf")

        if os.path.exists(md_path):
            full_paths.append(md_path)
            parse_markdown_to_docx(md_path, docx_path)
            parse_markdown_to_pdf(md_path, pdf_path)

    # Combined Document
    if full_paths:
        combined_docx = os.path.join(docs_dir, "SSF_AttnRes_Complete_Documentation.docx")
        combined_pdf = os.path.join(docs_dir, "SSF_AttnRes_Complete_Documentation.pdf")
        parse_markdown_to_docx(full_paths, combined_docx)
        parse_markdown_to_pdf(full_paths, combined_pdf)

if __name__ == "__main__":
    main()
